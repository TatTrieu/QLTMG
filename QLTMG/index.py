import os
from flask import render_template, request, redirect, url_for, session
from flask_login import login_user, logout_user, login_required, current_user
import dao
from QLTMG import app, login, db
from decorator import anonymous_required
from models import UserRole, Student, ClassRoom, Gender, HealthRecord, Regulation, Receipt, User
import admin
from datetime import datetime
import hashlib
from flask_mail import Mail, Message
from random import randint
from flask import Blueprint, send_file
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dateutil.relativedelta import relativedelta
import cloudinary
import cloudinary.uploader

# Lấy cấu hình từ Biến môi trường (Environment Variables)
cloudinary.config(
    cloud_name="dobee9dlb",
    api_key="158475847337887",
    api_secret="mKOhHPjUCnq8HczpUop3mfNFwaw"
)

login.login_view = 'login_process'

@login.user_loader
def load_user(user_id):
    return dao.get_user_by_id(user_id)


# --- 1. TRANG ĐĂNG NHẬP (Route /login) ---
@app.route("/login", methods=["GET", "POST"])
@anonymous_required  # Nếu đã đăng nhập rồi thì không cho vào trang này nữa (đá về dashboard)
def login_process():
    err_msg = None

    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")

        user = dao.auth_user(username, password)

        if user:
            login_user(user)
            # Đăng nhập thành công -> Chuyển về Dashboard (Trang chủ)
            # next_page xử lý trường hợp user đang định vào trang khác thì bị bắt đăng nhập
            next_page = request.args.get("next")
            return redirect(next_page if next_page else "/")
        else:
            err_msg = "Tài khoản hoặc mật khẩu không chính xác!"

    return render_template("login.html", err_msg=err_msg)


@app.route("/admin-login", methods=["POST"])
def admin_login_process():
    username = request.form.get("username")
    password = request.form.get("password")

    # 1. Xác thực user từ DB
    user = dao.auth_user(username, password)

    # 2. Kiểm tra: Phải có user VÀ user đó phải là ADMIN
    if user and user.role == UserRole.ADMIN:
        login_user(user)
        return redirect("/admin")
    else:
        err_msg = "Tài khoản không đúng hoặc không có quyền Admin!"

        # QUAN TRỌNG: Phải return về giao diện đăng nhập kèm lỗi
        # Sử dụng admin.index_view.render để giữ nguyên giao diện/menu của trang Admin
        return admin.index_view.render('admin/index.html', err_msg=err_msg)
# --- 2. TRANG ĐĂNG XUẤT ---
@app.route("/logout")
def logout_process():
    logout_user() # Xóa session của user hiện tại
    return redirect(url_for('login_process')) # Chuyển hướng ngay lập tức về trang đăng nhập


@app.route("/register", methods=['GET', 'POST'])
@login_required
def register_process():
    # Chỉ Admin mới được tạo tài khoản
    if current_user.role.name != 'ADMIN':
        return redirect(url_for('index', err_msg="Bạn không có quyền tạo tài khoản!"))

    err_msg = ""
    msg = ""

    if request.method == 'POST':
        name = request.form.get("name")
        username = request.form.get("username")
        password = request.form.get('password')
        confirm = request.form.get('confirm')
        email = request.form.get("email")

        # 1. XỬ LÝ UPLOAD AVATAR
        avatar_url = None  # Mặc định là None để hàm add_user tự lấy ảnh mặc định

        if 'avatar' in request.files:
            file = request.files['avatar']
            if file and file.filename != '':
                try:
                    # Upload lên Cloudinary
                    upload_result = cloudinary.uploader.upload(file)
                    avatar_url = upload_result['secure_url']
                except Exception as ex:
                    print(f"Lỗi upload ảnh: {ex}")
                    err_msg = "Lỗi: Không thể tải ảnh lên hệ thống!"
                    return render_template('register.html', err_msg=err_msg)

        # 2. KIỂM TRA MẬT KHẨU
        if password == confirm:
            # Gọi hàm add_user (Nếu avatar_url là None thì hàm DAO sẽ tự gán ảnh mặc định)
            if dao.add_user(name=name,
                            username=username,
                            password=password,
                            email=email,
                            avatar=avatar_url):  # Truyền link ảnh (hoặc None)

                # Thành công -> Về trang chủ
                return redirect(url_for('index', msg="Tạo tài khoản thành công!"))
            else:
                err_msg = "Lỗi! Tên đăng nhập hoặc Email đã tồn tại."
        else:
            err_msg = "Mật khẩu xác nhận không khớp!"

    return render_template('register.html', err_msg=err_msg, msg=msg)


@app.route("/")
@login_required
def index():
    # 1. Thông tin cơ bản trường (Thống kê)
    count_student = Student.query.filter(Student.active == True).count()
    count_class = ClassRoom.query.count()

    # 2. Lấy danh sách thông báo
    notifications = dao.load_notifications()

    return render_template("index.html",
                           count_student=count_student,
                           count_class=count_class,
                           notifications=notifications)


# --- XỬ LÝ THÊM THÔNG BÁO ---
@app.route("/notifications/add", methods=["POST"])
@login_required
def add_notification_process():
    # Kiểm tra quyền Admin
    if current_user.role.name != 'ADMIN':
        return redirect(url_for('index', err_msg="Truy cập bị từ chối!"))

    title = request.form.get('title')
    content = request.form.get('content')

    # [CẬP NHẬT] Truyền thêm current_user.id vào hàm DAO
    if dao.add_notification(title, content, user_id=current_user.id):
        return redirect(url_for('index'))
    else:
        return redirect(url_for('index', err_msg="Lỗi hệ thống!"))


# --- XỬ LÝ XÓA THÔNG BÁO ---
@app.route("/notifications/delete/<int:id>")
@login_required
def delete_notification(id):
    if current_user.role.name != 'ADMIN':
        return redirect(url_for('index'))

    dao.delete_notification(id)
    return redirect(url_for('index'))

# --- 2. TRANG THỐNG KÊ (STATS) ---
@app.route("/stats")
@login_required
def stats_process():
    # 1. Lấy NGÀY từ tham số (Thay vì lấy tháng)
    date_str = request.args.get('date')

    # Nếu không chọn, mặc định lấy Hôm nay
    if not date_str:
        date_str = datetime.now().strftime('%Y-%m-%d')

    # 2. Xử lý Phân quyền Lớp (Giữ nguyên code cũ)
    class_id = request.args.get('class_id')
    current_class_obj = None

    if current_user.role == UserRole.TEACHER:
        my_class = dao.get_class_by_teacher(current_user.id)
        if my_class:
            class_id = my_class.id
            current_class_obj = my_class
    else:
        if class_id and class_id != 'all':
            class_id = int(class_id)
            current_class_obj = ClassRoom.query.get(class_id)
        else:
            class_id = None

    # 3. Gọi DAO lấy dữ liệu (Truyền ngày cụ thể vào)
    stats_data = dao.get_dashboard_data(date_str, class_id)

    # Lấy cài đặt và danh sách lớp
    school_settings = dao.get_settings()
    classes = dao.load_classes()

    return render_template("stats.html",
                           data=stats_data,
                           current_date=date_str,  # Truyền biến ngày sang View
                           classes=classes,
                           current_class_obj=current_class_obj,
                           settings=school_settings)


@app.route("/students/add", methods=["POST"])
@login_required
def add_student_process():
    if current_user.role != UserRole.ADMIN:
        return redirect(url_for('students', err_msg="Bạn không có quyền thực hiện chức năng này!"))

    # 1. Lấy dữ liệu văn bản
    name = request.form.get("name")
    birth_date = request.form.get("birth_date")
    gender = request.form.get("gender")
    parent_name = request.form.get("parent_name")
    phone = request.form.get("phone")
    class_id = request.form.get("class_id")

    # 2. XỬ LÝ UPLOAD ẢNH (Thay vì lấy URL text)
    avatar_url = None
    if 'avatar' in request.files:
        file = request.files['avatar']
        if file and file.filename != '':
            try:
                # Upload lên Cloudinary
                upload_result = cloudinary.uploader.upload(file)
                avatar_url = upload_result['secure_url']
            except Exception as ex:
                print(f"Lỗi upload ảnh: {ex}")
                return redirect(url_for('students', err_msg="Lỗi: Không thể tải ảnh lên hệ thống!"))

    # 3. Validate Số điện thoại
    if phone:
        if not phone.isdigit():
             return redirect(url_for('students', err_msg="Lỗi: Số điện thoại chỉ được chứa ký tự số!"))
        if len(phone) < 9 or len(phone) > 11:
             return redirect(url_for('students', err_msg="Lỗi: Số điện thoại phải từ 9-11 số!"))

    # 4. Gọi DAO để lưu
    success, message = dao.add_student(name, birth_date, gender, parent_name, phone, class_id, avatar_url, current_user.id)

    if success:
        return redirect(url_for('students', msg=message))
    else:
        return redirect(url_for('students', err_msg=message))


# CẬP NHẬT LẠI ROUTE /students ĐỂ HIỂN THỊ THÔNG BÁO
@app.route("/students")
@login_required
def students():
    kw = request.args.get('kw')

    # --- LOGIC PHÂN QUYỀN LỚP ---
    class_id = request.args.get('class_id')  # Mặc định lấy từ URL (cho Admin)

    # Nếu là Giáo viên -> Tự động lấy lớp của họ
    if current_user.role.name == 'TEACHER':
        my_class = dao.get_class_by_teacher(current_user.id)
        if my_class:
            class_id = my_class.id  # Gán cứng class_id
        else:
            class_id = -1  # Giáo viên chưa được phân lớp -> Không hiện gì cả
    # -----------------------------

    msg = request.args.get('msg')
    err_msg = request.args.get('err_msg')

    # Các hàm dưới vẫn dùng biến class_id đã được xử lý ở trên
    list_students = dao.load_students(kw=kw, class_id=class_id)
    list_classes = dao.load_classes()
    system_settings = dao.get_settings()

    current_class = None
    if class_id and class_id != -1:
        current_class = dao.get_class_by_id(class_id)

    return render_template("students.html",
                           students=list_students,
                           classes=list_classes,
                           current_class_id=str(class_id) if class_id else None,
                           # Chuyển sang string để so sánh trong HTML
                           msg=msg,
                           err_msg=err_msg,
                           settings=system_settings,
                           current_class=current_class)

# --- XỬ LÝ XÓA ---
@app.route("/students/delete/<int:student_id>")
@login_required
def delete_student_process(student_id):
    # Kiểm tra quyền Admin (Nếu cần chặt chẽ)
    if current_user.role != UserRole.ADMIN:
        return redirect(url_for('students', err_msg="Bạn không có quyền xóa!"))

    if dao.delete_student(student_id):
        return redirect(url_for('students', msg="Đã xóa hồ sơ thành công!"))
    else:
        return redirect(url_for('students', err_msg="Lỗi khi xóa hồ sơ!"))

# --- XỬ LÝ CẬP NHẬT ---
@app.route("/students/update", methods=["POST"])
@login_required
def update_student_process():
    # 1. Lấy dữ liệu
    student_id = request.form.get("student_id")
    name = request.form.get("name")
    birth_date = request.form.get("birth_date")
    gender = request.form.get("gender")
    parent_name = request.form.get("parent_name")
    phone = request.form.get("phone")
    class_id = request.form.get("class_id")

    # 2. XỬ LÝ UPLOAD ẢNH (Update)
    avatar_url = None
    if 'avatar' in request.files:
        file = request.files['avatar']
        if file and file.filename != '':
            try:
                upload_result = cloudinary.uploader.upload(file)
                avatar_url = upload_result['secure_url']
            except Exception as ex:
                print(f"Lỗi upload ảnh: {ex}")
                return redirect(url_for('students', err_msg="Lỗi: Không thể tải ảnh lên hệ thống!"))

    # 3. Validate SĐT
    if phone:
        if not phone.isdigit():
            return redirect(url_for('students', err_msg="Lỗi: Số điện thoại chỉ được chứa ký tự số!"))
        if len(phone) < 9 or len(phone) > 11:
            return redirect(url_for('students', err_msg="Lỗi: Số điện thoại phải từ 9 đến 11 số!"))

    # 4. Gọi DAO cập nhật
    result = dao.update_student(student_id, name, birth_date, gender, parent_name, phone, class_id, avatar_url)

    # Xử lý kết quả trả về từ DAO (Tuple hoặc Boolean)
    if isinstance(result, tuple):
        success, message = result
    else:
        success, message = result, "Cập nhật thành công!" if result else "Lỗi cập nhật!"

    if success:
        return redirect(url_for('students', msg=message))
    else:
        return redirect(url_for('students', err_msg=message))


@app.route("/health")
@login_required
def health_process():
    kw = request.args.get('kw')

    # --- LOGIC PHÂN QUYỀN LỚP (Tương tự trên) ---
    class_id = request.args.get('class_id')

    if current_user.role.name == 'TEACHER':
        my_class = dao.get_class_by_teacher(current_user.id)
        if my_class:
            class_id = my_class.id
        else:
            class_id = -1
    # --------------------------------------------

    stats = dao.get_health_list_with_stats(kw=kw, class_id=class_id)
    classes = dao.load_classes()

    current_class = None
    if class_id and class_id != -1:
        current_class = dao.get_class_by_id(class_id)

    return render_template("health.html",
                           stats=stats,
                           classes=classes,
                           current_class_id=str(class_id) if class_id else None,
                           current_kw=kw,
                           current_class=current_class)


@app.route("/health/update", methods=["POST"])
@login_required
def update_health_process():
    student_id = request.form.get("student_id")

    # Lấy dữ liệu dạng chuỗi từ form
    height_str = request.form.get("height")
    weight_str = request.form.get("weight")
    temp_str = request.form.get("temperature")
    note = request.form.get("note")

    # --- [VALIDATION] KIỂM TRA SỐ ÂM ---
    try:
        # Chuyển đổi sang số thực để kiểm tra
        h = float(height_str) if height_str else 0
        w = float(weight_str) if weight_str else 0
        t = float(temp_str) if temp_str else 0

        if h < 0 or w < 0 or t < 0:
            return redirect(
                url_for('health_process', err_msg="Lỗi: Chiều cao, Cân nặng, Nhiệt độ không được nhập số âm!"))

        # Kiểm tra nhiệt độ hợp lý (Ví dụ không thể > 45 độ hoặc < 30 độ quá vô lý - Tùy chọn)
        if t > 45 or (t > 0 and t < 30):
            return redirect(url_for('health_process', err_msg="Lỗi: Nhiệt độ không hợp lý!"))

    except ValueError:
        return redirect(url_for('health_process', err_msg="Lỗi: Vui lòng nhập đúng định dạng số!"))
    # -----------------------------------

    if dao.add_new_health_checkup(student_id, height_str, weight_str, temp_str, note):
        return redirect(url_for('health_process', msg="Đã cập nhật số liệu mới thành công!"))
    else:
        return redirect(url_for('health_process', err_msg="Lỗi khi lưu dữ liệu!"))


@app.route("/tuition", methods=['GET'])
@login_required
def tuition():
    # --- [MỚI] Lấy thông báo từ URL (do hàm update gửi về) ---
    msg = request.args.get('msg')
    err_msg = request.args.get('err_msg')
    # --------------------------------------------------------

    # 1. XỬ LÝ DANH SÁCH THÁNG
    existing_months_query = db.session.query(Receipt.month).distinct().all()
    month_list = [m[0] for m in existing_months_query]

    try:
        month_list.sort(key=lambda x: datetime.strptime(x, '%m/%Y'), reverse=True)
    except:
        pass

    # 2. XỬ LÝ THÁNG ĐANG CHỌN
    month_str = request.args.get('month')
    class_id = request.args.get('class_id')

    if not month_str:
        month_str = datetime.now().strftime('%m/%Y')

    if month_str not in month_list:
        month_list.insert(0, month_str)

    try:
        curr_date = datetime.strptime(month_str, '%m/%Y')
        prev_month = (curr_date - relativedelta(months=1)).strftime('%m/%Y')
        next_month = (curr_date + relativedelta(months=1)).strftime('%m/%Y')
    except:
        curr_date = datetime.now()
        month_str = curr_date.strftime('%m/%Y')
        prev_month = (curr_date - relativedelta(months=1)).strftime('%m/%Y')
        next_month = (curr_date + relativedelta(months=1)).strftime('%m/%Y')

    # 3. XỬ LÝ PHÂN QUYỀN
    current_class_obj = None
    if current_user.role == UserRole.TEACHER:
        my_class = ClassRoom.query.filter_by(teacher_id=current_user.id).first()
        if my_class:
            class_id = my_class.id
            current_class_obj = my_class
        else:
            class_id = -1
    else:
        if class_id and class_id != 'all':
            current_class_obj = ClassRoom.query.get(class_id)

    # [AUTO UPDATE] Tự động tính lại tiền ăn theo điểm danh
    if month_str:
         dao.auto_update_tuition_from_attendance(month_str, class_id)

    # 4. LẤY DỮ LIỆU ĐỂ HIỂN THỊ
    reg_tuition = Regulation.query.filter_by(key='BASE_TUITION').first()
    reg_meal = Regulation.query.filter_by(key='MEAL_PRICE').first()
    base_price = reg_tuition.value if reg_tuition else 0
    meal_price = reg_meal.value if reg_meal else 0

    query = Student.query.filter(Student.active == True)
    if class_id and class_id != 'all' and class_id != -1:
        query = query.filter_by(class_id=class_id)
    elif class_id == -1:
        query = query.filter(Student.id == -1)
    students = query.all()

    data_list = []
    total_summary = {'meal_total': 0, 'discount': 0, 'total_due': 0, 'paid': 0, 'debt': 0}

    for s in students:
        receipt = Receipt.query.filter_by(student_id=s.id, month=month_str).first()
        item = {'student_id': s.id, 'name': s.name, 'month': month_str,
                'base_tuition': base_price, 'meal_price': meal_price}

        if receipt:
            item.update({'meal_days': receipt.meal_days, 'meal_total': receipt.meal_total,
                         'discount': receipt.discount, 'total_due': receipt.total_due,
                         'paid_amount': receipt.paid_amount, 'is_saved': True})
        else:
            default_days = 22
            meal_total_calc = default_days * meal_price
            item.update({'meal_days': default_days, 'meal_total': meal_total_calc,
                         'discount': 0, 'total_due': base_price + meal_total_calc,
                         'paid_amount': 0, 'is_saved': False})

        item['debt'] = item['total_due'] - item['paid_amount']

        total_summary['meal_total'] += item['meal_total']
        total_summary['discount'] += item['discount']
        total_summary['total_due'] += item['total_due']
        total_summary['paid'] += item['paid_amount']
        total_summary['debt'] += item['debt']

        data_list.append(item)

    classes = ClassRoom.query.all()

    return render_template("tuition.html",
                           data_list=data_list,
                           classes=classes,
                           total=total_summary,
                           current_month=month_str,
                           next_month=next_month,
                           prev_month=prev_month,
                           month_list=month_list,
                           current_class_id=str(class_id) if class_id else None,
                           current_class_obj=current_class_obj,
                           msg=msg,          # <--- Đã thêm biến này
                           err_msg=err_msg)  # <--- Đã thêm biến này

# --- ROUTE MỚI: KHỞI TẠO DỮ LIỆU THÁNG ---
@app.route("/tuition/init", methods=['POST'])
@login_required
def init_tuition_data():
    month = request.form.get('month')
    class_id = request.form.get('class_id')

    # 1. Lấy Quy định giá tiền
    reg_tuition = Regulation.query.filter_by(key='BASE_TUITION').first()
    reg_meal = Regulation.query.filter_by(key='MEAL_PRICE').first()
    base_price = reg_tuition.value if reg_tuition else 0
    meal_price = reg_meal.value if reg_meal else 0

    # 2. Lấy danh sách học sinh cần tạo
    query = Student.query.filter(Student.active == True)
    if class_id and class_id != 'all' and class_id != '-1':
        query = query.filter_by(class_id=class_id)
    students = query.all()

    count = 0
    for s in students:
        # Kiểm tra xem đã có hóa đơn chưa
        existing = Receipt.query.filter_by(student_id=s.id, month=month).first()

        if not existing:
            # Tạo mới với số liệu mặc định
            default_days = 22
            meal_total = default_days * meal_price
            total_due = base_price + meal_total

            new_receipt = Receipt(
                student_id=s.id,
                month=month,
                meal_days=default_days,
                base_tuition=base_price,
                meal_total=meal_total,
                discount=0,
                total_due=total_due,
                paid_amount=0,
                status=False,
                user_id=current_user.id
            )
            db.session.add(new_receipt)
            count += 1

    db.session.commit()

    return redirect(
        url_for('tuition', month=month, class_id=class_id, msg=f"Đã khởi tạo dữ liệu cho {count} học sinh!"))


# --- ROUTE XỬ LÝ LƯU (KHI BẤM NÚT "LƯU TRẠNG THÁI") ---
@app.route("/tuition/update-single", methods=['POST'])
@login_required
def update_single_tuition():
    # 1. Lấy dữ liệu cơ bản
    student_id = request.form.get('student_id')
    month = request.form.get('month')

    # Lấy class_id để khi redirect về đúng trang lớp đó (UX tốt hơn)
    # Vì form modal của bạn có thể không gửi class_id, ta sẽ lấy student để tìm class nếu cần
    # Tuy nhiên, đơn giản nhất là redirect về tuition với month, user sẽ tự lọc lại.

    # --- HÀM PHỤ: Xử lý chuỗi rỗng thành 0.0 ---
    def safe_float(value):
        if not value or value.strip() == '':
            return 0.0
        try:
            return float(value)
        except ValueError:
            return 0.0

    # -------------------------------------------

    # 2. Lấy giá trị từ Form (Dùng safe_float để không bị crash)
    meal_days = safe_float(request.form.get('meal_days'))
    discount = safe_float(request.form.get('discount'))
    paid_amount = safe_float(request.form.get('paid_amount'))

    # 3. Lấy đơn giá quy định
    reg_tuition = Regulation.query.filter_by(key='BASE_TUITION').first()
    reg_meal = Regulation.query.filter_by(key='MEAL_PRICE').first()
    base_price = reg_tuition.value if reg_tuition else 0
    meal_price = reg_meal.value if reg_meal else 0

    # 4. Tính toán TỔNG CẦN THU (Total Due) trước
    meal_total = meal_days * meal_price
    total_due = base_price + meal_total - discount

    # --- [QUAN TRỌNG] LOGIC KIỂM TRA HỢP LỆ ---

    # Kiểm tra số âm
    if paid_amount < 0:
        return redirect(url_for('tuition', month=month, err_msg="Lỗi: Số tiền đã đóng không được âm!"))
    if discount < 0:
        return redirect(url_for('tuition', month=month, err_msg="Lỗi: Số tiền miễn giảm không được âm!"))
    if meal_days < 0:
        return redirect(url_for('tuition', month=month, err_msg="Lỗi: Số ngày ăn không được âm!"))

    # Kiểm tra đóng thừa tiền (Đóng > Phải thu)
    # Dùng round để tránh lỗi so sánh số thực nhỏ (vd: 0.0000001)
    if round(paid_amount) > round(total_due):
        str_paid = "{:,.0f}".format(paid_amount)
        str_due = "{:,.0f}".format(total_due)
        return redirect(url_for('tuition', month=month,
                                err_msg=f"Lỗi: Số tiền đóng ({str_paid}) không được lớn hơn số phải thu ({str_due})!"))

    # ------------------------------------------

    # 5. Lưu vào Database
    is_finished = (total_due - paid_amount) <= 0

    receipt = Receipt.query.filter_by(student_id=student_id, month=month).first()

    if receipt:
        # Update
        receipt.meal_days = meal_days
        receipt.meal_total = meal_total
        receipt.discount = discount
        receipt.total_due = total_due
        receipt.paid_amount = paid_amount
        receipt.status = is_finished
        receipt.user_id = current_user.id
    else:
        # Create
        new_receipt = Receipt(
            student_id=student_id,
            month=month,
            meal_days=meal_days,
            base_tuition=base_price,
            meal_total=meal_total,
            discount=discount,
            total_due=total_due,
            paid_amount=paid_amount,
            status=is_finished,
            user_id=current_user.id
        )
        db.session.add(new_receipt)

    db.session.commit()

    # Load lại trang và thông báo thành công
    return redirect(url_for('tuition', month=month, msg="Cập nhật thành công!"))


@app.route("/settings", methods=["GET", "POST"])
@login_required
def settings_process():
    # Kiểm tra quyền Admin
    if current_user.role.name != 'ADMIN':
        return redirect(url_for('index', err_msg="Truy cập bị từ chối!"))

    msg = None
    err_msg = None

    if request.method == "POST":
        data = {
            'MAX_STUDENT': request.form.get('max_student'),
            'BASE_TUITION': request.form.get('base_tuition'),
            'MEAL_PRICE': request.form.get('meal_price')
        }

        # [CẬP NHẬT] Truyền thêm current_user.id để lưu vết người sửa
        success, message = dao.update_settings(data, user_id=current_user.id)

        if success:
            msg = message
        else:
            err_msg = message

    current_settings = dao.get_settings()
    return render_template("settings.html", settings=current_settings, msg=msg, err_msg=err_msg)


@app.route("/profile", methods=['GET', 'POST'])
@login_required
def profile():
    err_msg = ""
    msg = ""

    if request.method == 'POST':
        # 1. Lấy thông tin cơ bản
        name = request.form.get('name')
        email = request.form.get('email')

        # 2. XỬ LÝ UPLOAD ẢNH (Thay thế hoàn toàn logic nhập URL cũ)
        avatar_url = None  # Mặc định là None (nghĩa là không đổi ảnh)

        if 'avatar' in request.files:
            file = request.files['avatar']
            # Chỉ xử lý nếu người dùng có chọn file
            if file and file.filename != '':
                try:
                    # Upload lên Cloudinary
                    upload_result = cloudinary.uploader.upload(file)
                    # Lấy đường dẫn ảnh trả về
                    avatar_url = upload_result['secure_url']
                except Exception as e:
                    print(f"Lỗi Cloudinary: {e}")
                    # QUAN TRỌNG: Nếu lỗi upload thì dừng ngay, không lưu DB
                    return render_template('profile.html',
                                           err_msg="Lỗi: Không thể tải ảnh lên (Vui lòng kiểm tra kết nối mạng hoặc API Key)!",
                                           msg="")

        # 3. Lấy thông tin mật khẩu
        old_password = request.form.get('old_password')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')

        # 4. LOGIC KIỂM TRA BẢO MẬT (Xác minh mật khẩu cũ)
        # Người dùng phải nhập mật khẩu cũ KHI:
        # - Có thay đổi Email
        # - Hoặc Có nhập mật khẩu mới
        is_sensitive_change = False

        if email and email.strip() != (current_user.email or ""):
            is_sensitive_change = True

        if new_password:
            is_sensitive_change = True
            if new_password != confirm_password:
                return render_template('profile.html', err_msg="Mật khẩu mới và xác nhận không khớp!", msg="")

        if is_sensitive_change:
            if not old_password:
                return render_template('profile.html',
                                       err_msg="Vui lòng nhập mật khẩu hiện tại để xác nhận thay đổi quan trọng!",
                                       msg="")

            # Kiểm tra mật khẩu cũ có đúng không
            hashed_old_pass = hashlib.md5(old_password.strip().encode('utf-8')).hexdigest()
            if hashed_old_pass != current_user.password:
                return render_template('profile.html', err_msg="Mật khẩu hiện tại không đúng!", msg="")

        # 5. GỌI DAO ĐỂ LƯU VÀO DATABASE
        # Truyền avatar_url (Nếu là None thì DAO sẽ giữ nguyên ảnh cũ)
        if dao.update_user_profile(current_user.id, name, email, avatar_url, new_password):
            msg = "Cập nhật hồ sơ thành công!"
        else:
            err_msg = "Lỗi! Email này có thể đã được sử dụng bởi người khác."

    return render_template('profile.html', err_msg=err_msg, msg=msg)

# --- 1. CẤU HÌNH MAIL (Đặt đoạn này sau khi khởi tạo app) ---
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'nguyentattrieuu2005@gmail.com' # <--- Thay Email của bạn vào đây
app.config['MAIL_PASSWORD'] = 'jyou wsoc hroj qhqx'  # <--- Thay Mật khẩu ứng dụng vào đây
mail = Mail(app)


# GIAI ĐOẠN 1: NHẬP EMAIL & GỬI OTP
@app.route("/forgot-password", methods=['GET', 'POST'])
def forgot_password():
    err_msg = ""
    if request.method == 'POST':
        email = request.form.get('email')
        user = dao.get_user_by_email(email)

        if user:
            # 1. Tạo OTP ngẫu nhiên 6 số
            otp = randint(100000, 999999)

            # 2. Lưu OTP và Email vào Session tạm thời
            session['otp'] = str(otp)
            session['reset_email'] = email

            # 3. Gửi Email
            try:
                msg = Message('Mã xác thực đổi mật khẩu - Mầm Non TBT',
                              sender='noreply@mamnon.com',
                              recipients=[email])
                msg.body = f"Chào bạn,\nMã xác thực (OTP) của bạn là: {otp}\nMã này có hiệu lực để đổi mật khẩu."
                mail.send(msg)

                # Chuyển sang trang nhập OTP
                return redirect(url_for('verify_otp'))
            except Exception as ex:
                err_msg = f"Lỗi gửi mail: {str(ex)}"
        else:
            err_msg = "Email này chưa được đăng ký trong hệ thống!"

    return render_template('forgot_password.html', err_msg=err_msg)


# GIAI ĐOẠN 2: XÁC MINH OTP
@app.route("/verify-otp", methods=['GET', 'POST'])
def verify_otp():
    err_msg = ""
    if request.method == 'POST':
        otp_input = request.form.get('otp')
        otp_session = session.get('otp')

        if otp_session and otp_input == otp_session:
            # OTP đúng -> Chuyển sang trang đặt lại mật khẩu
            return redirect(url_for('reset_new_password'))
        else:
            err_msg = "Mã OTP không chính xác!"

    return render_template('verify_otp.html', err_msg=err_msg)


# GIAI ĐOẠN 3: ĐẶT MẬT KHẨU MỚI
@app.route("/reset-password", methods=['GET', 'POST'])
def reset_new_password():
    err_msg = ""
    email = session.get('reset_email')  # Lấy email từ session

    if not email:  # Nếu không có email trong session (truy cập trái phép)
        return redirect(url_for('login_process'))

    if request.method == 'POST':
        password = request.form.get('password')
        confirm = request.form.get('confirm')

        if password == confirm:
            if dao.update_password_by_email(email, password):
                # Xóa session sau khi thành công
                session.pop('otp', None)
                session.pop('reset_email', None)
                return redirect(url_for('login_process', msg="Đổi mật khẩu thành công! Vui lòng đăng nhập."))
            else:
                err_msg = "Lỗi hệ thống khi cập nhật mật khẩu."
        else:
            err_msg = "Mật khẩu xác nhận không khớp!"

    return render_template('reset_password.html', err_msg=err_msg)



# --- Route 2: Xử lý xuất file Word (Logic chính) ---
@app.route('/export/student-list-word', methods=['POST'])
def export_student_list_word():
    # 1. Lấy dữ liệu từ Form
    class_id = request.form.get('class_id')
    note_content = request.form.get('note', '')

    # 2. Kiểm tra lớp học
    selected_class = ClassRoom.query.get(class_id)
    if not selected_class:
        return "Lỗi: Không tìm thấy lớp!", 404

    # 3. Lấy danh sách học sinh
    students = Student.query.filter_by(class_id=class_id).all()

    # 4. Khởi tạo file Word
    document = Document()

    # -- Tiêu đề file --
    # Để giống mẫu: "Biểu mẫu danh sách lớp:" hoặc Tiêu đề lớn
    document.add_heading(f'DANH SÁCH HỌC SINH - LỚP {selected_class.name.upper()}', 0)

    # -- Chèn Ghi chú (nếu có) --
    if note_content:
        document.add_heading('Ghi chú:', level=2)
        p = document.add_paragraph(note_content)
        p.italic = True

    document.add_paragraph('')  # Xuống dòng

    # -- Tạo bảng với 6 CỘT đúng theo mẫu ảnh --
    # Cột: STT | Họ tên trẻ | Ngày sinh | Giới tính | Phụ huynh | SĐT liên hệ
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # -- Thiết lập Tiêu đề bảng --
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Họ tên trẻ'
    hdr_cells[2].text = 'Ngày sinh'
    hdr_cells[3].text = 'Giới tính'
    hdr_cells[4].text = 'Phụ huynh'
    hdr_cells[5].text = 'SĐT liên hệ'

    # -- Đổ dữ liệu --
    for index, student in enumerate(students, start=1):
        row_cells = table.add_row().cells

        # 1. STT
        row_cells[0].text = str(index)

        # 2. Họ tên
        row_cells[1].text = student.name

        # 3. Ngày sinh (Format dd/mm/yyyy)
        if student.birth_date:
            row_cells[2].text = student.birth_date.strftime('%d/%m/%Y')
        else:
            row_cells[2].text = ""

        # 4. Giới tính (Chuyển Enum thành Tiếng Việt)
        # Kiểm tra Enum Gender từ models.py
        if student.gender == Gender.MALE:
            row_cells[3].text = "Nam"
        elif student.gender == Gender.FEMALE:
            row_cells[3].text = "Nữ"
        else:
            row_cells[3].text = ""

        # 5. Phụ huynh
        row_cells[4].text = student.parent_name if student.parent_name else ""

        # 6. SĐT liên hệ
        row_cells[5].text = student.phone if student.phone else ""

    # -- Thêm dòng quy định giống mẫu ảnh (Optional) --
    document.add_paragraph('')
    # Lấy thông số sĩ số từ setting (nếu cần) hoặc text cứng như ảnh
    # document.add_paragraph('Quy định: Mỗi lớp tối đa 25 trẻ.')

    # 5. Xuất file
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    safe_classname = selected_class.name.replace(' ', '_')
    filename = f"DanhSach_{safe_classname}.docx"

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/export/health-list-word', methods=['POST'])
def export_health_list_word():
    # 1. Lấy dữ liệu từ Form
    class_id = request.form.get('class_id')
    note_content = request.form.get('note', '')

    selected_class = ClassRoom.query.get(class_id)
    if not selected_class:
        return "Lỗi: Không tìm thấy lớp!", 404

    students = Student.query.filter_by(class_id=class_id).all()

    # 2. Khởi tạo file Word
    document = Document()
    document.add_heading(f'BÁO CÁO SỨC KHỎE - LỚP {selected_class.name.upper()}', 0)

    # Hiển thị ghi chú
    if note_content:
        document.add_heading('Ghi chú:', level=2)
        p = document.add_paragraph(note_content)
        p.italic = True

    document.add_paragraph(f'Ngày cập nhật: {datetime.now().strftime("%d/%m/%Y")}')
    document.add_paragraph('')  # Xuống dòng cho thoáng

    # --- TẠO BẢNG 6 CỘT ---
    # Cấu trúc: STT | Tên | Chiều cao | Cân nặng | Nhiệt độ | Ghi chú
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Thiết lập tiêu đề cột
    hdr = table.rows[0].cells
    hdr[0].text = 'STT'
    hdr[1].text = 'Họ tên trẻ'
    hdr[2].text = 'Chiều cao (cm)'  # <-- Đã thêm cột này
    hdr[3].text = 'Cân nặng (kg)'
    hdr[4].text = 'Nhiệt độ'
    hdr[5].text = 'Ghi chú'

    # Chỉnh độ rộng cột (Optional - giúp bảng đẹp hơn)
    # table.columns[0].width = Cm(1.0) # STT
    # table.columns[1].width = Cm(5.0) # Tên
    # table.columns[2].width = Cm(2.5) # Chiều cao
    # table.columns[3].width = Cm(2.5) # Cân nặng
    # table.columns[4].width = Cm(2.0) # Nhiệt độ
    # table.columns[5].width = Cm(4.0) # Ghi chú

    # Đổ dữ liệu
    for idx, s in enumerate(students, 1):
        row = table.add_row().cells

        # 1. STT
        row[0].text = str(idx)

        # 2. Tên
        row[1].text = s.name

        # Lấy dữ liệu sức khỏe mới nhất
        hr = HealthRecord.query.filter_by(student_id=s.id) \
            .order_by(HealthRecord.created_date.desc()) \
            .first()

        if hr:
            # 3. Chiều cao
            row[2].text = str(hr.height) if hr.height else ""

            # 4. Cân nặng
            row[3].text = str(hr.weight) if hr.weight else ""

            # 5. Nhiệt độ
            row[4].text = str(hr.temperature) if hr.temperature else ""

            # 6. Ghi chú
            row[5].text = hr.note if hr.note else ""
        else:
            # Nếu chưa có dữ liệu thì để trống tất cả
            row[2].text = ""
            row[3].text = ""
            row[4].text = ""
            row[5].text = ""

    # Xuất file
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    filename = f"SucKhoe_{selected_class.name.replace(' ', '_')}.docx"

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/export/tuition-list-word', methods=['POST'])
@login_required
def export_tuition_list_word():
    # 1. Lấy dữ liệu từ Form (Modal gửi lên)
    month = request.form.get('month')
    class_id = request.form.get('class_id')
    note_content = request.form.get('note', '')

    # 2. Lấy thông tin lớp
    selected_class = ClassRoom.query.get(class_id)
    class_name = selected_class.name if selected_class else "Tất cả"

    # 3. Lấy quy định giá (để tính toán nếu chưa có hóa đơn)
    reg_tuition = Regulation.query.filter_by(key='BASE_TUITION').first()
    reg_meal = Regulation.query.filter_by(key='MEAL_PRICE').first()
    base_price = reg_tuition.value if reg_tuition else 0
    meal_price = reg_meal.value if reg_meal else 0

    # 4. Lấy danh sách học sinh
    query = Student.query.filter(Student.active == True)
    if class_id and class_id != 'all' and class_id != '-1':
        query = query.filter_by(class_id=class_id)
    students = query.all()

    # 5. Khởi tạo file Word
    document = Document()

    # Tiêu đề
    p_title = document.add_heading(f'BÁO CÁO THU HỌC PHÍ - THÁNG {month}', 0)
    p_title.alignment = 1  # Căn giữa

    document.add_paragraph(f'Lớp: {class_name}')
    document.add_paragraph(f'Ngày xuất báo cáo: {datetime.now().strftime("%d/%m/%Y")}')

    # Ghi chú user nhập
    if note_content:
        document.add_heading('Ghi chú:', level=3)
        p = document.add_paragraph(note_content)
        p.italic = True

    document.add_paragraph('')  # Xuống dòng

    # 6. Tạo bảng dữ liệu (8 CỘT)
    # STT | Họ tên | Học phí CB | Tiền ăn | Miễn giảm | Phải thu | Đã thu | Còn nợ
    table = document.add_table(rows=1, cols=8)
    table.style = 'Table Grid'

    # Header
    hdr = table.rows[0].cells
    headers = ['STT', 'Họ tên trẻ', 'Học phí CB', 'Tiền ăn', 'Miễn giảm', 'Phải thu', 'Đã thu', 'Còn nợ']
    for i, text in enumerate(headers):
        hdr[i].text = text
        # In đậm header
        run = hdr[i].paragraphs[0].runs
        if run: run[0].font.bold = True

    # Biến tổng cộng
    sum_total_due = 0
    sum_paid = 0
    sum_debt = 0

    # 7. Đổ dữ liệu
    for idx, s in enumerate(students, 1):
        row = table.add_row().cells

        # Lấy thông tin hóa đơn
        receipt = Receipt.query.filter_by(student_id=s.id, month=month).first()

        # Tính toán dữ liệu
        if receipt:
            r_base = base_price  # Hoặc receipt.base_tuition nếu muốn lấy lịch sử
            r_meal = receipt.meal_total
            r_discount = receipt.discount
            r_due = receipt.total_due
            r_paid = receipt.paid_amount
        else:
            # Nếu chưa có hóa đơn -> Tính mặc định (22 ngày)
            default_days = 22
            r_base = base_price
            r_meal = default_days * meal_price
            r_discount = 0
            r_due = r_base + r_meal
            r_paid = 0

        r_debt = r_due - r_paid

        # Cộng tổng
        sum_total_due += r_due
        sum_paid += r_paid
        sum_debt += r_debt

        # Định dạng tiền tệ (Thêm dấu phẩy: 1,500,000)
        def fmt(money):
            return "{:,.0f}".format(money)

        # Ghi vào ô
        row[0].text = str(idx)
        row[1].text = s.name
        row[2].text = fmt(r_base)
        row[3].text = fmt(r_meal)
        row[4].text = fmt(r_discount)
        row[5].text = fmt(r_due)
        row[6].text = fmt(r_paid)
        row[7].text = fmt(r_debt)

    # 8. Thêm dòng TỔNG CỘNG cuối bảng
    row_sum = table.add_row().cells
    row_sum[0].merge(row_sum[4])  # Merge từ ô 0 đến ô 4
    row_sum[0].text = "TỔNG CỘNG:"
    row_sum[0].paragraphs[0].runs[0].font.bold = True
    row_sum[0].paragraphs[0].alignment = 2  # Căn phải

    row_sum[5].text = "{:,.0f}".format(sum_total_due)
    row_sum[5].paragraphs[0].runs[0].font.bold = True

    row_sum[6].text = "{:,.0f}".format(sum_paid)
    row_sum[6].paragraphs[0].runs[0].font.bold = True

    row_sum[7].text = "{:,.0f}".format(sum_debt)
    row_sum[7].paragraphs[0].runs[0].font.bold = True
    row_sum[7].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)

    # 9. Xuất file
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    filename = f"HocPhi_{month.replace('/', '_')}_{class_name.replace(' ', '_')}.docx"

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route("/attendance", methods=['GET', 'POST'])
@login_required
def attendance():
    # 1. Xử lý logic chọn lớp (Giống trang học phí)
    class_id = request.args.get('class_id')

    # Nếu là Giáo viên -> Tự gán lớp của họ
    current_class_obj = None
    if current_user.role.name == 'TEACHER':
        my_class = dao.get_class_by_teacher(current_user.id)
        if my_class:
            class_id = my_class.id
            current_class_obj = my_class
    else:
        # Nếu là Admin mà có chọn lớp
        if class_id and class_id != 'all':
            current_class_obj = dao.get_class_by_id(class_id)

    # 2. Lấy ngày chọn (Mặc định là hôm nay)
    date_str = request.args.get('date')
    if not date_str:
        date_str = datetime.now().strftime('%Y-%m-%d')

    # 3. Lấy danh sách điểm danh
    student_list = []
    if class_id:  # Chỉ hiện khi đã xác định được lớp
        student_list = dao.get_attendance_list(class_id, date_str)

    classes = dao.load_classes()  # Để Admin chọn lớp

    return render_template('attendance.html',
                           students=student_list,
                           classes=classes,
                           current_class_id=str(class_id) if class_id else None,
                           current_class_obj=current_class_obj,
                           current_date=date_str)


@app.route("/attendance/save", methods=['POST'])
@login_required
def save_attendance_process():
    # Lấy dữ liệu từ form
    class_id = request.form.get('class_id')
    date_str = request.form.get('date')

    # Form sẽ gửi về danh sách dạng: status_1, note_1, status_2, note_2... (với số là id học sinh)
    # Chúng ta cần duyệt qua tất cả key trong form để tìm dữ liệu
    for key in request.form:
        if key.startswith('status_'):
            student_id = key.split('_')[1]  # Lấy ID học sinh từ tên input
            status = request.form.get(key)
            note = request.form.get(f'note_{student_id}')

            # Gọi hàm lưu
            dao.save_attendance(student_id, date_str, status, note)

    return redirect(url_for('attendance', class_id=class_id, date=date_str, msg="Đã lưu điểm danh thành công!"))


# --- CHẠY APP ---
if __name__ == "__main__":
    app.run(debug=True, port=5000)